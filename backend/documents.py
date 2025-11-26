"""
Documents Generator Module
Contains: Document models, NDA/MSA/SOW generation, MSA template generation
"""
from fastapi import APIRouter, HTTPException, Depends
from pydantic import BaseModel, Field
from typing import Dict, Any, Optional
import uuid
from datetime import datetime, timezone
import logging
from pathlib import Path
from docx import Document
from io import BytesIO
import base64
from mailmerge import MailMerge
import tempfile
import shutil

from login import User, get_current_user

logger = logging.getLogger(__name__)

# Router
router = APIRouter(prefix="/api", tags=["documents"])

# Database reference - will be set from server.py
db = None

# Root directory reference - will be set from server.py
ROOT_DIR = None

def set_db(database):
    global db
    db = database

def set_root_dir(root_dir):
    global ROOT_DIR
    ROOT_DIR = root_dir

# ============== MODELS ==============

class DocumentRequest(BaseModel):
    template_type: str  # nda/msa/sow
    engagement_model: str  # t&m/dedicated/fixed_bid
    variables: Dict[str, Any]

class MSADocumentRequest(BaseModel):
    date: str
    company_name: str
    customer_company_address: str
    title: str
    name: str

class DocumentResponse(BaseModel):
    id: str
    template_type: str
    doc_base64: str
    message: Optional[str] = None

# ============== HELPER FUNCTIONS ==============

def create_doc_template(template_type: str, variables: Dict[str, Any]) -> BytesIO:
    """Create a document from template with mail merge fields"""
    doc = Document()

    # Set document title
    title = doc.add_heading(template_type.upper(), 0)

    if template_type == "nda":
        doc.add_heading('Non-Disclosure Agreement', level=1)
        doc.add_paragraph(f"This Agreement is made on {variables.get('date', 'DATE')}")
        doc.add_paragraph(f"Between: {variables.get('party1_name', 'PARTY_1_NAME')} ('{variables.get('party1_short', 'Party 1')}')")
        doc.add_paragraph(f"And: {variables.get('party2_name', 'PARTY_2_NAME')} ('{variables.get('party2_short', 'Party 2')}')")
        doc.add_paragraph('\n')
        doc.add_heading('1. Confidential Information', level=2)
        doc.add_paragraph(variables.get('confidential_info', 'The parties agree to protect confidential information shared during business discussions.'))
        doc.add_heading('2. Obligations', level=2)
        doc.add_paragraph(variables.get('obligations', 'Both parties agree not to disclose confidential information to third parties without prior written consent.'))
        doc.add_heading('3. Term', level=2)
        doc.add_paragraph(f"This agreement shall remain in effect for {variables.get('term', 'TERM_PERIOD')}.")

    elif template_type == "msa":
        doc.add_heading('Master Service Agreement', level=1)
        doc.add_paragraph(f"Effective Date: {variables.get('effective_date', 'EFFECTIVE_DATE')}")
        doc.add_paragraph(f"Client: {variables.get('client_name', 'CLIENT_NAME')}")
        doc.add_paragraph(f"Service Provider: {variables.get('provider_name', 'PROVIDER_NAME')}")
        doc.add_paragraph('\n')
        doc.add_heading('1. Services', level=2)
        doc.add_paragraph(variables.get('services_desc', 'Services to be provided as per Statement of Work.'))
        doc.add_heading('2. Payment Terms', level=2)
        doc.add_paragraph(variables.get('payment_terms', 'Payment terms as specified in individual SOWs.'))
        doc.add_heading('3. Term and Termination', level=2)
        doc.add_paragraph(f"Initial term: {variables.get('term', 'TERM_PERIOD')}")

    elif template_type == "sow":
        doc.add_heading('Statement of Work', level=1)
        doc.add_paragraph(f"Project: {variables.get('project_name', 'PROJECT_NAME')}")
        doc.add_paragraph(f"Client: {variables.get('client_name', 'CLIENT_NAME')}")
        doc.add_paragraph(f"Engagement Model: {variables.get('engagement_model', 'ENGAGEMENT_MODEL')}")
        doc.add_paragraph('\n')
        doc.add_heading('1. Scope of Work', level=2)
        doc.add_paragraph(variables.get('scope', 'Project scope and deliverables.'))
        doc.add_heading('2. Timeline', level=2)
        doc.add_paragraph(f"Duration: {variables.get('duration', 'DURATION')}")
        doc.add_paragraph(f"Start Date: {variables.get('start_date', 'START_DATE')}")
        doc.add_heading('3. Resources', level=2)
        doc.add_paragraph(variables.get('resources', 'Team composition and resource allocation.'))
        doc.add_heading('4. Cost', level=2)
        doc.add_paragraph(f"Total Cost: ${variables.get('total_cost', '0.00')}")
        doc.add_paragraph(variables.get('cost_breakdown', 'Cost breakdown as per engagement model.'))

    # Add signatures section
    doc.add_page_break()
    doc.add_heading('Signatures', level=2)
    doc.add_paragraph('\n\n')
    doc.add_paragraph(f"{variables.get('signatory1_name', 'SIGNATORY_1')}")
    doc.add_paragraph(f"{variables.get('signatory1_title', 'Title')}")
    doc.add_paragraph(f"Date: {variables.get('signature_date', '_____________')}")
    doc.add_paragraph('\n\n')
    doc.add_paragraph(f"{variables.get('signatory2_name', 'SIGNATORY_2')}")
    doc.add_paragraph(f"{variables.get('signatory2_title', 'Title')}")
    doc.add_paragraph(f"Date: {variables.get('signature_date', '_____________')}")

    # Save to BytesIO
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

def format_date_string(date_str: str) -> str:
    """Convert YYYY-MM-DD to '17th October 2025' format"""
    try:
        if not date_str or not date_str.strip():
            return "<To Be Filled>"

        date_obj = datetime.strptime(date_str, "%Y-%m-%d")
        day = date_obj.day

        # Add ordinal suffix (st, nd, rd, th)
        if 4 <= day <= 20 or 24 <= day <= 30:
            suffix = "th"
        else:
            suffix = ["st", "nd", "rd"][day % 10 - 1]

        formatted = date_obj.strftime(f"{day}{suffix} %B %Y")
        return formatted
    except:
        return "<To Be Filled>"

def sanitize_field(value: Any) -> str:
    """Replace empty/whitespace-only values with '<To Be Filled>'"""
    if value is None:
        return "<To Be Filled>"

    str_value = str(value).strip()
    if not str_value:
        return "<To Be Filled>"

    return str_value

# ============== DOCUMENT GENERATOR ROUTES ==============

@router.post("/documents/generate", response_model=DocumentResponse)
async def generate_document(request: DocumentRequest, current_user: User = Depends(get_current_user)):
    doc_id = str(uuid.uuid4())

    # Generate document
    doc_stream = create_doc_template(request.template_type, request.variables)

    # Convert to base64
    doc_base64 = base64.b64encode(doc_stream.read()).decode('utf-8')

    # Save document metadata
    doc_record = {
        "id": doc_id,
        "template_type": request.template_type,
        "engagement_model": request.engagement_model,
        "variables": request.variables,
        "doc_url": "generated",
        "doc_base64": doc_base64,
        "status": "generated",
        "created_by": current_user.id,
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    await db.documents.insert_one(doc_record)

    return DocumentResponse(
        id=doc_id,
        template_type=request.template_type,
        doc_base64=doc_base64,
        message="Document generated successfully"
    )

@router.get("/documents")
async def get_documents(current_user: User = Depends(get_current_user)):
    docs = await db.documents.find({"created_by": current_user.id}, {"_id": 0}).to_list(1000)
    return docs

# ============== MSA DOCUMENT GENERATION ==============

@router.post("/documents/msa/generate")
async def generate_msa_document(request: MSADocumentRequest, current_user: User = Depends(get_current_user)):
    """Generate MSA document from template and return DOCX"""
    doc_id = str(uuid.uuid4())

    try:
        # Path to template
        template_path = ROOT_DIR.parent / "Zuci MSA Template.docx"

        if not template_path.exists():
            raise HTTPException(status_code=404, detail="MSA template not found")

        # Create temporary directory for file operations
        temp_dir = Path(tempfile.mkdtemp())

        try:
            filled_docx_path = temp_dir / f"MSA_{doc_id}.docx"

            # Open template and merge fields
            document = MailMerge(str(template_path))

            # Prepare merge fields
            merge_fields = {
                "date": format_date_string(request.date),
                "company_name": sanitize_field(request.company_name),
                "customer_company_address": sanitize_field(request.customer_company_address),
                "title": sanitize_field(request.title),
                "name": sanitize_field(request.name),
            }

            # Get all merge fields from template and fill missing ones
            template_fields = document.get_merge_fields()
            for field in template_fields:
                if field not in merge_fields:
                    merge_fields[field] = "<To Be Filled>"

            # Merge and save
            document.merge(**merge_fields)
            document.write(str(filled_docx_path))

            # Read DOCX and encode to base64
            with open(filled_docx_path, 'rb') as f:
                docx_content = f.read()
                docx_base64 = base64.b64encode(docx_content).decode('utf-8')

            doc_record = {
                "id": doc_id,
                "template_type": "msa",
                "engagement_model": "custom",
                "variables": request.model_dump(),
                "doc_url": "generated",
                "status": "generated",
                "created_by": current_user.id,
                "created_at": datetime.now(timezone.utc).isoformat()
            }

            await db.documents.insert_one(doc_record)

            return {
                "id": doc_id,
                "doc_base64": docx_base64,
                "template_type": "msa",
                "message": "Document generated successfully"
            }

        finally:
            # Cleanup temp directory
            shutil.rmtree(temp_dir, ignore_errors=True)

    except Exception as e:
        logger.error(f"Error generating MSA document: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to generate document: {str(e)}")

@router.get("/documents/msa/preview")
async def preview_msa_template(current_user: User = Depends(get_current_user)):
    """Get merge fields from MSA template"""
    try:
        template_path = ROOT_DIR.parent / "Zuci MSA Template.docx"

        if not template_path.exists():
            raise HTTPException(status_code=404, detail="MSA template not found")

        document = MailMerge(str(template_path))
        merge_fields = document.get_merge_fields()

        return {
            "merge_fields": list(merge_fields),
            "template_name": "Zuci MSA Template.docx"
        }
    except Exception as e:
        logger.error(f"Error reading template: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to read template: {str(e)}")
